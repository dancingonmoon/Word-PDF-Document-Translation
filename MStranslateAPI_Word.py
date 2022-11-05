import json
import uuid

import docx
import requests
import math
import re
import time
import fitz
# import copy

# define the function, input the text and ouput the translation:
def MStranslation_API(
    text,
    lang_in="en",
    lang_out="zh-Hans",
    subscription_key="pls input your MStranslation API key",
):
    """
    To translate the pre-defined text:
    args:
        text: text or text list, to be translated;
        lang_in : language input;
        lang_out: language output, language list is accepted, for example:['ru','zh-Hans']
        subscription_key: Microsoft Translation API Key;
    out:
        trans_text: the tranlated text list (when single language output), or translated text dictionary (when multiple language output); when in dictionary, the key is the language names, the value is the list of translated text
        response: API output, as backup, in case of overseeing the errors;
    """

    # Add your subscription key and endpoint
    subscription_key = subscription_key
    endpoint = "https://api.cognitive.microsofttranslator.com"  
    # Add your location, also known as region. The default is global.
    # This is required if using a Cognitive Services resource.
    location = "eastasia"
    path = "/translate"
    constructed_url = endpoint + path

    params = {"api-version": "3.0", "from": lang_in, "to": lang_out}

    headers = {
        "Ocp-Apim-Subscription-Key": subscription_key,
        "Ocp-Apim-Subscription-Region": location,
        "Content-type": "application/json",
        "X-ClientTraceId": str(uuid.uuid4()),
    }

    # You can pass more than one object in body.
    body = []

    if isinstance(text, list):
        for txt in text:
            body.append({"text": txt})

    if isinstance(text, str):
        body.append({"text": text})

    # print(body)
    # to turn lang_out into list, when lang_out is with one language:
    if isinstance(lang_out, str):
        lang_out = [lang_out]

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    # response_json = json.dumps(response, sort_keys=True, ensure_ascii=False, indent=4, separators=(',', ': '))
    # print(response_json)

    trans_text = []

    try:
        if isinstance(lang_out, list) and len(lang_out) > 1:
            trans_text = {}
            for i, lang in enumerate(lang_out):
                tmp = []
                for r in response:
                    tmp.append(r["translations"][i]["text"])
                trans_text[lang] = tmp
        else:
            for r in response:
                trans_text.append(r["translations"][0]["text"])
    except:
        print(response)

    return trans_text, response

def MStranslation_dynamicDictionary_API(
    text,
    dynamic_dict=False,
    lang_in="en",
    lang_out="zh-Hans",
    subscription_key="pls input your MStranslation API key",
):
    """
    Microsoft Translation API with Dynamic dictionary:
    args:
        text: text or text list to be translated;
        dynamic_dict: dynamic dictionary, comprises of specialized vocabulary, production name, personal names etc., that has conventional translated phrases for example: {'莫言':'Mr.Moyan'};
        lang_in : language input;
        lang_out: language output, list is accepted to add multiple language, for example:['ru','zh-Hans']
        subscription_key: Microsoft Translation API key;
    out:
        trans_text: the text list (when single language output), or text dictionary (when multiple language), dictionary shall have language name as the key, and translated text list as value
        response: API output, as backup, in case of overseeing the errors;
    """

    # Add your subscription key and endpoint
    subscription_key = subscription_key
    endpoint = "https://api.cognitive.microsofttranslator.com"  
    # Add your location, also known as region. The default is global.
    # This is required if using a Cognitive Services resource.
    location = "eastasia"
    path = "/translate"
    constructed_url = endpoint + path

    params = {"api-version": "3.0", "from": lang_in, "to": lang_out}

    headers = {
        "Ocp-Apim-Subscription-Key": subscription_key,
        "Ocp-Apim-Subscription-Region": location,
        "Content-type": "application/json",
        "X-ClientTraceId": str(uuid.uuid4()),
    }

    # You can pass more than one object in body.
    body = []

    if isinstance(text, list):
        for txt in text:
            if isinstance(dynamic_dict, dict):
                for key in dynamic_dict.keys():
                    # if txt comprises of dyanamic_dict key, then key shall be replaced by value.
                    sub_txt = (
                        '<mstrans:dictionary translation="'
                        + dynamic_dict[key]
                        + '">'
                        + key
                        + "</mstrans:dictionary>"
                    )
                    txt = re.sub(key, sub_txt, txt)
                body.append({"text": txt})
            elif dynamic_dict == False or dynamic_dict == "":
                body.append({"text": txt})
            else:
                print("Neither False nor Dictionary dynamic_dict is !")

    if isinstance(text, str):  # when text is one single text string
        if isinstance(dynamic_dict, dict):
            for key in dynamic_dict.keys():
                # if txt comprises of dyanamic_dict key, then key shall be replaced by value.
                sub_txt = (
                    '<mstrans:dictionary translation="'
                    + dynamic_dict[key]
                    + '">'
                    + key
                    + "</mstrans:dictionary>"
                )
                text = re.sub(key, sub_txt, text)
            body.append({"text": text})
        elif dynamic_dict == False or dynamic_dict == "":
            body.append({"text": text})
        else:
            print("Neither False nor Dictionary dynamic_dict is !")

    # when lang_out is with single language, to turn lang_out into list:
    if isinstance(lang_out, str):
        lang_out = [lang_out]

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    # response_json = json.dumps(response, sort_keys=True, ensure_ascii=False, indent=4, separators=(',', ': '))
    # print(response_json)

    trans_text = []

    try:
        if isinstance(lang_out, list) and len(lang_out) > 1:
            trans_text = {}
            for i, lang in enumerate(lang_out):
                tmp = []
                for r in response:
                    tmp.append(r["translations"][i]["text"])
                trans_text[lang] = tmp
        else:
            for r in response:
                trans_text.append(r["translations"][0]["text"])
    except:
        print(response)

    return trans_text, response



# to define fucntion, for each 'Run' in each paragraph object (stalled in memory), perceive whether the "Run" contains inlined picture/drawing, only non-null text shall be applied and subsititued:
def paragraph_runs_replace(
    paragraph,
    text,
):
    """
    To replace the text of (inside) each run of each paragraph, with inlined shape/picture/drawings untouched;
    1. for the paragraph object, go through each run, get the ratio between run.text and paragraph.text, then apply such ratio upon translated text, so as to distribute the translated text upon each of "Run" and get the distriubted run.text start and stop pointer in the form of tuple;
    2. for each run, to replace each of run.text , when run.text is non-null.
    Note: the fuction revoke the library of python-docx, which operate on document object. Document object is stalled in memory; simple copy or deep copy can't work on object.
    args:
        paragraph: paragraph object, could be from document, or table, or header,footer etc.;
        text: replaced text, single string, not list;
    out:
        paragraph: paragraph object that complete the operation;
    """
    text_len = len(text)  # to get the length of text string;
    source_len = len(paragraph.text)

    run_attr = {}
    pointer = 0
    for i, r in enumerate(paragraph.runs):
        if source_len == 0:
            run_attr[i] = 0
        else:
            len_distrib = math.ceil(len(r.text) / source_len * text_len)  # round up to integer
            run_attr[i] = (pointer, pointer + len_distrib)
            pointer = pointer + len_distrib
    # print(run_attr)
    
    # go though each run, replace run.text if non-null:
    if len(paragraph.runs) != 0:
        for i, r in enumerate(paragraph.runs):
            if r.text != "":
                # replace run.text, with style and font unchanged:
                r.text = text[run_attr[i][0] : run_attr[i][1]]

    return paragraph

def Word_MStranslate(
    doc,
    dynamic_dict=False,
    lang_in='zh-Hans',
    lang_out='en',    
    filename=False,
    subscription_key="pls input your MStranslation API key",
):
    """
    To work on paragraph, table, header and footer, and replace those text with translated text, and style and font remain; if dynamic dictionary is in need, text with dictionary keys shall be replaced with dictionary value;
    A: paragraph:
    1. generate the dictionary for the paragraphs structure, and the text;
    2. Set up dictionary, with the key has the value that is the index No. inside the test list;
    2. invoke the translation API, to generate translted_text list;
    3. for each paragraph, to replace the text with translated text, with style and font unchanged;
    B: Table :
    1. generate the dictionary for the table structure, and the text
    2. Set up dictionary, with the key has the value that is the index No. inside the test list;
    2. invoke the translation API, to generate translted_text list;
    3. for each table, to replace the text with translated text, with style and font unchanged;
    C: header and footer translation;
    args:
        doc: document object generated from python-docx;
        lang_in : language input;
        lang_out: language output, in single language;
        dynamic_dict: dynamic dictionary, comprises of specialized vocabulary, production name, personal names etc., that has conventional translated phrases for example: {'莫言':'Mr.Moyan'};
        filename: whether to save into specified file path; default value:  False, otherwise, the file path
        subscription_key: Microsoft API key;
    out:
        doc: the document object after the operation;
        when filename!=False, to save document into specified path;
        if error comes out, API outputs the response with error code;
    """
    # Paragraphs Translation:
    # -----------------------------------
    # generate text list to be translated
    text_dict = {}
    for i, para in enumerate(doc.paragraphs):
        text_dict[(i)] = para.text
    # print(text_dict)
    text = list(text_dict.values())
    # Set up dictionary, with the key has the value that is the index No. inside the test list;
    textindex_dict = {}
    for key in text_dict.keys():
        textindex_dict[key] = text.index(text_dict[key])

    # generate the translated text list
    # invoke API, to deicde which API is gona to be used ,subject to the dynamic dictionary;
    if dynamic_dict == False:
        trans_text, _ = MStranslation_API(
            text, lang_in=lang_in, lang_out=lang_out, subscription_key=subscription_key
        )
    elif isinstance(dynamic_dict, dict):
        trans_text, _ = MStranslation_dynamicDictionary_API(
            text,
            dynamic_dict=dynamic_dict,
            lang_in=lang_in,
            lang_out=lang_out,
            subscription_key=subscription_key,
        )
    else:
        print("Neither False nor Dictionary dynamic_dict is !")
        
    # for each paragraph:
    for i, para in enumerate(doc.paragraphs):
        para_trans_text = trans_text[textindex_dict[(i)]]
        paragraph_runs_replace(para, para_trans_text)

    # -------------------------------------------

    # Table translation:
    # -------------------------------------
    # get the dictionary for the table structure, and the text
    text_dict = {}
    for t, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                for p, para in enumerate(cell.paragraphs):
                    text_dict[(t, r, c, p)] = para.text

    # print('text字典len:{}'.format(len(text)))
    text = list(text_dict.values())
    # Set up dictionary, with the key has the value that is the index No. inside the test list;
    textindex_dict = {}
    for key in text_dict.keys():
        textindex_dict[key] = text.index(text_dict[key])
        
    # invoke API, to deicde which API is gona to be used ,subject to the dynamic dictionary;
    if dynamic_dict == False:
        trans_text, _ = MStranslation_API(
            text, lang_in=lang_in, lang_out=lang_out, subscription_key=subscription_key
        )
    elif isinstance(dynamic_dict, dict):
        trans_text, _ = MStranslation_dynamicDictionary_API(
            text,
            dynamic_dict=dynamic_dict,
            lang_in=lang_in,
            lang_out=lang_out,
            subscription_key=subscription_key,
        )
    else:
        print("Neither False nor Dictionary dynamic_dict is !")

    # for each paragraph of table structure:
    for t, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                for p, para in enumerate(cell.paragraphs):

                    para_trans_text = trans_text[textindex_dict[(t, r, c, p)]]
                    paragraph_runs_replace(para, para_trans_text)

    # ----------------------------------
    # header, footer , translation:
    # ----------------------------------
    # to get all section.header.paragraphs和secction.footer.paragraphs structure in dictionary and the text
    text_dict = {}
    for s, section in enumerate(doc.sections):
        for p, para in enumerate(section.header.paragraphs):  # each section has only one header;
            text_dict[(s, "header", p)] = para.text
        for p, para in enumerate(section.footer.paragraphs):  # each section only has one footer;
            if para.text.isdigit() == False and para.text != "":  # footer that has dynamic page No, let it be
                text_dict[(s, "footer", p)] = para.text
    # print(text_dict)
    text = list(text_dict.values())
    # Set up dictionary, with the key has the value that is the index No. inside the test list;   
    textindex_dict = {}
    for key in text_dict.keys():
        textindex_dict[key] = text.index(text_dict[key])

    # invoke API, to deicde which API is gona to be used ,subject to the dynamic dictionary;
    if dynamic_dict == False:
        trans_text, _ = MStranslation_API(
            text, lang_in=lang_in, lang_out=lang_out, subscription_key=subscription_key
        )
    elif isinstance(dynamic_dict, dict):
        trans_text, _ = MStranslation_dynamicDictionary_API(
            text,
            dynamic_dict=dynamic_dict,
            lang_in=lang_in,
            lang_out=lang_out,
            subscription_key=subscription_key,
        )
    else:
        print("Neither False nor Dictionary dynamic_dict is !")
        
    # for each paragraph of all section.header.paragraphs:
    for s, section in enumerate(doc.sections):
        for p, para in enumerate(section.header.paragraphs):  # each section only has one header;
            para_trans_text = trans_text[textindex_dict[(s, "header", p)]]
            paragraph_runs_replace(para, para_trans_text)
        for p, para in enumerate(section.footer.paragraphs):  # each section only has one header;
            if para.text.isdigit() == False and para.text != "":  # footer with dynamic pages No, leave it unchanged;
                para_trans_text = trans_text[textindex_dict[(s, "footer", p)]]
                paragraph_runs_replace(para, para_trans_text)
    # ---------------------------------
    # save,output:
    if filename != False:
        doc.save(filename)
    return doc


def PDF_MStranslate(
    input_path,
    output_path,
    lang_in="en",
    lang_out="zh-Hans",
    dynamic_dict=False,
    subscription_key="pls input your MS_translation API Key",
    image2txt=False,
    txtbox_borderColor="gray",
    out_font="cjk",
):
    """
    1. Read each image, including inlined image, and write into a new PDF;
    2. Read each block text, send into MS_Translation_API, to get trans_text list;
    3. Set up dictionary with block/line/span sequence No. as the key, add the value with corresponding span_text, fontsize,color, length, etc.;
    4. To divide trans_text according to the distribution of spen_ext;
    5. To insert span_trans_text into block/line/span , starting form its original position, with same fontsize,color, etc.;
    6. To save new PDF into out_path.
    Args:
        input_path: Source PDF file path;
        output_path:the file path where to save output PDF;
        lang_in: the source PDF language;
        lang_out: the tranlated language of PDF;
        dynamic_dict: dynamic dictionary used by MS_translation API, including special vocabulary, production name, persona name,etc.,for exampe:{'莫言':'Mr.Moyan'};default=False,to suggest no dynamic dictionary is in need;
        subscription_key: MS_translation API Key;
        image2txt: whether to repsent the image with text summary; default=False, to suggest putting original image;
        txtbox_borderColor: the text box boarder color, the Default is "Gray", when txtbox_borderColor=False时,to suggest no text boarder is in need;
        out_font: the fontname that is to display in output PDF, for chinese/japanese/korea, could be "cjk", for english, could 'helv', or anyone you like;
    Out:
        to produce new PDF, with the above operation, that is to copy image, and translate text; and save the PDF, close the new and old PDF files.

    """

    t0 = time.time()
    # doc1 = fitz.open(sys.argv[1])
    doc1 = fitz.open(input_path)
    doc2 = fitz.open()
    pink = fitz.utils.getColor("pink")  # give out the color RGB tuple
    blue = fitz.utils.getColor("blue")
    # green = fitz.utils.getColor("green")
    if txtbox_borderColor != False:
        txtbox_borderColor = fitz.utils.getColor(txtbox_borderColor)
    else:
        txtbox_borderColor = None
    # gray = (0.9, 0.9, 0.9)

    # to insert new font, to display multiple language:
    font_cjk = fitz.Font(out_font)

    for page1 in doc1:
        # produce new PDF with same size
        page2 = doc2.new_page(
            -1, width=page1.mediabox_size[0], height=page1.mediabox_size[1]
        )
        # the text font we use
        # fontname= the customized name you like; but the font comes from fontbuffer
        page2.insert_font(fontname="cjk", fontbuffer=font_cjk.buffer)

        # draw rectangle:
        img = page2.new_shape()  # prepare /Contents object
        # calculate /CropBox & displacement
        disp = fitz.Rect(
            page1.cropbox_position, page1.cropbox_position
        )  # to get doc1 coordinate
        croprect = page1.rect + disp  # add to get the displacement;

        # draw original /CropBox rectangle
        img.draw_rect(croprect)
        img.finish(color=txtbox_borderColor, fill=None)

        # get image write into PDF, translate, put back into PDF:
        blocks = page1.get_text("dict")
        blocks_ = page1.get_text("blocks")
        # set up dictionary, to get span_text, fontsize, color, etc
        span_attr = {}
        text = []
        for b, block in enumerate(blocks["blocks"]):
            if block["type"] == 1:  # image
                rect = fitz.Rect(block["bbox"]) 
                if image2txt:
                    a = fitz.TEXT_ALIGN_CENTER
                    block_txt = blocks_[b][4]  # image description
                    rect += disp
                    img.draw_rect(rect)  # surround block rectangle
                    img.finish(width=0.3, color=pink)
                    img.insert_textbox(
                        rect, buffer=block_txt, fontsize=8, color=pink, align=a
                    )
                else:
                    page2.insert_image(
                        rect,
                        stream=block["image"],
                    )

            if block["type"] == 0:  # block为text
                block_txt = blocks_[b][4]
                # remove '\n'
                block_txt = re.sub("\n", "", block_txt)
                block_txt_len = len(block_txt)  
                text.append(block_txt)

                for l, line in enumerate(block["lines"]):
                    for s, span in enumerate(line["spans"]):
                        span_attr[(b, l, s, "size")] = span["size"]
                        span_attr[(b, l, s, "font")] = span["font"]
                        span_attr[(b, l, s, "color")] = span["color"]
                        span_attr[(b, l, s, "origin")] = span["origin"]
                        span_attr[(b, l, s, "text_index")] = len(text) - 1
                        span_attr[(b, l, s, "span_block_ratio")] = (
                            len(span["text"]) / block_txt_len
                        )

        try:
            trans_text, response = MStranslation_dynamicDictionary_API(
                text,
                dynamic_dict=dynamic_dict,
                lang_in=lang_in,
                lang_out=lang_out,
                subscription_key=subscription_key,
            )
        except:
            print(response)

        # distribute trans_text upon the span_text 

        for b, block in enumerate(blocks["blocks"]):
            if block["type"] == 0:  # text block
                rect = fitz.Rect(block["bbox"])
                # add dislacement of original /CropBox
                rect += disp

                img.draw_rect(rect)  # surround block rectangle
                a = fitz.TEXT_ALIGN_LEFT

                img.finish(width=0.3, color=txtbox_borderColor)

                pointer = 0
                for l, line in enumerate(block["lines"]):
                    for s, span in enumerate(line["spans"]):
                        span_transtxt_len = math.ceil(
                            len(trans_text[span_attr[(b, l, s, "text_index")]])
                            * span_attr[(b, l, s, "span_block_ratio")]
                        )
                        
                        span_attr[(b, l, s, "trans_text")] = trans_text[
                            span_attr[(b, l, s, "text_index")]
                        ][pointer : pointer + span_transtxt_len]
                        pointer = pointer + span_transtxt_len
                       

                        if rect.is_empty:  # do not rely on meaningful rects
                            print(
                                "skipping text of empty rect at ({}, {}) on page {}".format(
                                    rect.x0, rect.y0, page1.number
                                )
                            )
                        else:
                            # dict gives out the color in format of sRGB, to convert it to RGB, with float(0,1):
                            color = fitz.utils.sRGB_to_pdf(
                                span_attr[(b, l, s, "color")]
                            )
                            point = fitz.Point(
                                span_attr[(b, l, s, "origin")]
                            )  
                            
                            img.insert_text(
                                point=point,
                                buffer=span_attr[(b, l, s, "trans_text")],
                                fontname="cjk",
                                fontsize=span_attr[(b, l, s, "size")],
                                color=color,
                            )  

        img.commit()  # store /Contents of out page

    # save output file
    doc2.save(output_path, garbage=4, deflate=True, clean=True)
    doc1.close()
    doc2.close()  
    t1 = time.time()
    print("total time: {:.2f} sec".format((t1 - t0)))


